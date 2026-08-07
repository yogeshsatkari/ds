import json
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate

from data_to_docx.validate import TEMPLATE_PATH, validate_context

FONT_NAME = "Liberation Serif"
FONT_SIZE = Pt(12)
FIRST_COL_WIDTH = Inches(2.8)
DATE_COL_WIDTH = Inches(0.85)


def _set_cell_margins(cell, *, top=40, bottom=40, left=80, right=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tc_pr.append(margins)


def _set_cell_width(cell, width) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    cell_width = OxmlElement("w:tcW")
    cell_width.set(qn("w:w"), str(int(width.twips)))
    cell_width.set(qn("w:type"), "dxa")
    tc_pr.append(cell_width)


def _set_repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def _write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    font_size=FONT_SIZE,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = font_size
    r_fonts = run._element.rPr.rFonts
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)


def _format_date_header(date_str: str) -> str:
    for fmt in ("%d/%m/%y", "%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
        try:
            parsed = datetime.strptime(date_str.strip(), fmt)
            return f"{parsed.strftime('%Y-%m-')}\n{parsed.strftime('%d')}"
        except ValueError:
            continue
    return date_str


def _add_paragraph(subdoc, text: str, *, bold: bool = False) -> None:
    paragraph = subdoc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    r_fonts = run._element.rPr.rFonts
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)


def build_investigation_reports(doc: DocxTemplate, context: dict):
    if not context.get("include_investigation_reports"):
        return ""

    reports = context.get("investigation_reports") or []
    if not reports:
        return ""

    subdoc = doc.new_subdoc()
    for report in reports:
        title = report["title"].strip()
        if title and not title.endswith(":"):
            title = f"{title}:"
        _add_paragraph(subdoc, title, bold=True)

        for finding in report["findings"]:
            finding_text = finding.strip()
            if not finding_text:
                continue
            if not finding_text.startswith("•"):
                finding_text = f"• {finding_text}"
            _add_paragraph(subdoc, finding_text)

    return subdoc


def build_investigation_table(doc: DocxTemplate, context: dict):
    if not context.get("include_investigation_table"):
        return ""

    dates = context.get("investigation_dates") or []
    rows = context.get("investigation_rows") or []
    if not rows or not dates:
        return ""

    subdoc = doc.new_subdoc()
    table = subdoc.add_table(rows=1 + len(rows), cols=1 + len(dates))
    table.style = "Table Grid"
    table.autofit = False
    _set_repeat_header_row(table.rows[0])

    header = table.rows[0].cells
    _set_cell_width(header[0], FIRST_COL_WIDTH)
    _set_cell_margins(header[0])
    _write_cell(header[0], "Investigation", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for index, date in enumerate(dates, start=1):
        cell = header[index]
        _set_cell_width(cell, DATE_COL_WIDTH)
        _set_cell_margins(cell, left=40, right=40)
        _write_cell(
            cell,
            _format_date_header(date),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row_index, row in enumerate(rows, start=1):
        cells = table.rows[row_index].cells
        _set_cell_width(cells[0], FIRST_COL_WIDTH)
        _set_cell_margins(cells[0])
        _write_cell(cells[0], row["name"])

        for col_index, value in enumerate(row["values"], start=1):
            cell = cells[col_index]
            _set_cell_width(cell, DATE_COL_WIDTH)
            _set_cell_margins(cell, left=40, right=40)
            _write_cell(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER)

    return subdoc


def load_template_bytes(template_path: Path = TEMPLATE_PATH) -> bytes:
    if not template_path.is_file():
        raise FileNotFoundError(f"Missing Word template: {template_path}")
    return template_path.read_bytes()


def render_discharge_summary_bytes(context: dict, template_bytes: bytes) -> bytes:
    validate_context(context)
    doc = DocxTemplate(BytesIO(template_bytes))
    context = dict(context)
    context["investigation_table"] = build_investigation_table(doc, context)
    context["investigation_reports"] = build_investigation_reports(doc, context)
    doc.render(context)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_discharge_summary_from_template(
    context: dict,
    template_path: Path = TEMPLATE_PATH,
) -> bytes:
    return render_discharge_summary_bytes(context, load_template_bytes(template_path))


def render_discharge_summary(
    template_path: Path,
    context: dict,
    output_path: Path,
) -> Path:
    if not template_path.is_file():
        raise FileNotFoundError(f"Missing Word template: {template_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_bytes = render_discharge_summary_bytes(context, template_path.read_bytes())
    output_path.write_bytes(output_bytes)
    return output_path
